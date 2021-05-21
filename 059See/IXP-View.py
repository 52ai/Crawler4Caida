# coding:utf-8
"""
create on Apr 19, 2021 By Wenyan YU
Email: ieeflsyu@outlook.com

Function:

为进一步美化全球互联网交换中心数据分析报告，实现公众号自动化输出报告的功能
特对原报告进行完善，增加图绘制，和word报告自动生成功能


- - - - - - - - - - - - - - -历史记录 - - - - - - - - - - - - - -
create on May 29, 2020 By Wenyan YU
Function：

缘起新型互联网交换中心，感谢PeeringDB给了面试数据分析的基础，时间过去已近三年
今天再回过头来看PeeringDB，不禁感慨万千
我经常跟自己说，未来是数据的时代，数据资源的积累，数据处理方法论的探索，大数据技术的研究便显得非常重要

该程序的主要目的是再次探索分析PeeringDB的IX信息（https://www.peeringdb.com/api/ix）
{'id': 1,
'org_id': 2,
'name': 'Equinix Ashburn',
'name_long': 'Equinix Ashburn Exchange',
'city': 'Ashburn',
'country': 'US',
'region_continent': 'North America',
'media': 'Ethernet',
'notes': '',
'proto_unicast': True,
'proto_multicast': False,
'proto_ipv6': True,
'website': 'https://ix.equinix.com',
'url_stats': '',
'tech_email': 'support@equinix.com',
'tech_phone': '',
'policy_email': 'support@equinix.com',
'policy_phone': '',
'net_count': 324,
'created': '2010-07-29T00:00:00Z',
'updated': '2016-11-23T21:40:34Z',
'status': 'ok'}

可借助IX创建的时间以及国家、城市、大洲的信息进行相关分析，包括是否支持IPV6等

1）当前时间全球IXP总数的统计
2）全球IXP发展趋势
3）全球IXP信息按国家（地区）维度统计
4）全球IXP信息按大洲维度统计
5）我国IXP发展具体情况
6）全球IXP的IPV6支持情况
7）案例1：阿姆斯特丹交换中心（AMS-IX）数据统计分析(https://peeringdb.com/api/ix/26)
8）案例2：香港交换中心（HK-IX）数据统计分析
9）案例3：莫斯科交换中心（MSK-IX Moscow）数据统计分析

以上所有数据均可实时统计，可基于PeeringDB的数据一键生产当前时间【全球互联网交换中心的数据分析报告（PEERING DB）】

20210521 PEERING DB API研究
查询某个网络接入交换中心的数据，如Google 15169网的net_id 为433
API：https://www.peeringdb.com/api/netixlan?net_id=433

查询某个网络私有互联设施的数据，如Google 15169网的net_id 为433
API：https://www.peeringdb.com/api/netfac?net_id=433


"""
from urllib.request import urlopen
import json
from datetime import *
import time
from docx import Document
from docx.shared import Inches

import plotly.express as px
import random
import numpy as np


def generate_global_ixp_report(doc_file_path):
    """
    基于PEERING DB数据一键生成当前时间【全球互联网交换中心数据分析报告】
    :param doc_file_path:
    :return:
    """
    print(doc_file_path)
    html = urlopen(r'https://www.peeringdb.com/api/ix')
    html_json = json.loads(html.read())
    # print(html_json)
    ixp_cnt_year = {}  # 统计每年IXP的数量
    country_dict = {}  # 统计当前时间每个国家的IXP数量
    region_dict = {}  # 统计当前时间每个大洲的IXP数量

    ixp_cn = []  # 中国大陆的IXP
    ixp_hk = []  # 香港地区
    ixp_tw = []  # 台湾地区
    ixp_us = []  # 美国
    ixp_de = []  # 德国
    ixp_br = []  # 巴西
    ixp_ca = []  # 加拿大

    ipv6_on_cnt = 0  # 统计支撑ipv6的交换中心
    ipv6_off_cnt = 0  # 统计不支持ipv6的交换中心

    for item in html_json['data']:
        # 利用datetime处理时间字符串
        dt_create = datetime.strptime(item['created'], '%Y-%m-%dT%H:%M:%SZ')
        # print(dt_create)
        # 创建时间统计
        if dt_create.year not in ixp_cnt_year.keys():
            ixp_cnt_year[int(dt_create.year)] = 1
        else:
            ixp_cnt_year[int(dt_create.year)] = ixp_cnt_year[int(dt_create.year)] + 1
        # 按国家统计
        if item['country'] not in country_dict.keys():
            country_dict[item['country']] = 1
        else:
            country_dict[item['country']] = country_dict[item['country']] + 1
        # 按大洲统计
        if item['region_continent'] not in region_dict.keys():
            region_dict[item['region_continent']] = 1
        else:
            region_dict[item['region_continent']] = region_dict[item['region_continent']] + 1
        # 统计我国IXP相关情况
        temp = []
        if item['country'] == 'CN':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            ixp_cn.append(temp)

        temp = []
        if item['country'] == 'HK':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            ixp_hk.append(temp)

        temp = []
        if item['country'] == 'TW':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            ixp_tw.append(temp)

        temp = []
        if item['country'] == 'US':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            temp.append(item['city'])
            ixp_us.append(temp)

        temp = []
        if item['country'] == 'DE':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            temp.append(item['city'])
            ixp_de.append(temp)

        temp = []
        if item['country'] == 'BR':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            temp.append(item['city'])
            ixp_br.append(temp)

        temp = []
        if item['country'] == 'CA':
            temp.append(item['name'])
            temp.append(item['created'])
            temp.append(item['website'])
            temp.append(item['name_long'])
            temp.append(item['city'])
            ixp_ca.append(temp)

        # print(item['proto_ipv6'])
        if item['proto_ipv6']:
            ipv6_on_cnt += 1
        else:
            ipv6_off_cnt += 1

    # 新建一个文档
    document = Document()
    document.add_heading("- - - - - - -0)全球互联网交换中心数据分析报告- - ", level=1)
    temp_str = "报告生成时间：", str(datetime.now())
    document.add_paragraph(temp_str)
    document.add_paragraph("基础数据来源：https://www.peeringdb.com/")
    document.add_heading("- - - - - - -1)当前时间全球IXP总数统计- - - - - - - - - - - -", level=1)
    temp_str = "截止目前，全球共有", str(len(html_json['data'])), "个互联网交换中心"
    document.add_paragraph(temp_str)
    document.add_heading("- - - - - - -2)全球IXP发展趋势- - - - - - - - - - - -", level=1)
    ixp_cnt_year_list = []  # 存储统计列表
    for key in ixp_cnt_year.keys():
        ixp_cnt_year_list.append([key, ixp_cnt_year[key]])
    ixp_cnt_year_list.sort(key=lambda elem: int(elem[0]))
    # document.add_paragraph(ixp_cnt_year_list)
    # document.add_paragraph("全球互联网自2010年至今，每年新增的IXP数量:")
    # for item in ixp_cnt_year_list:
    #     temp_str = str(item[0]), "年新增IXP数量(个):", str(item[1])
    #     document.add_paragraph(temp_str)

    document.add_paragraph("全球互联网自2010年至今，每年总计的IXP数量:")
    years = []
    nums = []
    total_ixp = 0
    for item in ixp_cnt_year_list:
        total_ixp += int(item[1])
        temp_str = str(item[0]), "年总计IXP数量(个):", str(total_ixp)
        document.add_paragraph(temp_str)
        years.append(int(item[0]))
        nums.append(total_ixp)

    # print(ixp_cnt_year_list)
    fig = px.bar(x=years, y=nums, labels={'x': '年份', 'y': '全球IXP数量'})
    fig.write_image("bar_figure.png", engine="kaleido")
    document.add_picture('bar_figure.png', width=Inches(5.0))

    document.add_heading("- - - - - - -3)全球IXP信息按国家（地区）维度统计- - - - - - - - - - - -", level=1)
    country_dict_list = []  # 存储统计列表
    for key in country_dict.keys():
        country_dict_list.append([key, country_dict[key]])
    country_dict_list.sort(reverse=True, key=lambda elem: int(elem[1]))
    # document.add_paragraph(country_dict_list)
    temp_str = "全球范围内共有", str(len(country_dict_list)), "个国家（地区）部署了IXP"
    document.add_paragraph(temp_str)
    document.add_paragraph("按IXP数量降序排名，TOP 20信息如下：")
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'国家'
    hdr_cells[1].text = u'IXP数量'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True

    for item in country_dict_list[0:20]:
        # temp_str = str(item[0]), ":", str(item[1])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])

    document.add_heading("- - - - - - -4)全球IXP信息按大洲维度统计- - - - - - - - - - - -", level=1)
    region_dict_list = []  # 存储统计列表
    for key in region_dict.keys():
        region_dict_list.append([key, region_dict[key]])
    region_dict_list.sort(reverse=True, key=lambda elem: int(elem[1]))
    document.add_paragraph("按照IXP数量降序排名，各大洲IXP数量分布如下：")
    for item in region_dict_list:
        temp_str = str(item[0]), ":", str(item[1])
        document.add_paragraph(temp_str)

    document.add_heading("- - - - - - -5)我国IXP发展具体情况- - - - - - - - - - - -", level=1)
    temp_str = "我国大陆地区CN的IXP数量:", str(len(ixp_cn)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_cn:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]

    temp_str = "我国香港地区HK的IXP数量:", str(len(ixp_hk)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_hk:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]

    temp_str = "我国台湾地区TW的IXP数量:", str(len(ixp_tw)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_tw:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]

    temp_str = "美国的IXP数量:", str(len(ixp_us)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    hdr_cells[4].text = u'城市'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[4].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_us:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]

    temp_str = "德国的IXP数量:", str(len(ixp_de)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    hdr_cells[4].text = u'城市'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[4].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_de:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]

    temp_str = "巴西的IXP数量:", str(len(ixp_br)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    hdr_cells[4].text = u'城市'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[4].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_br:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]

    temp_str = "加拿大的IXP数量:", str(len(ixp_ca)), "，其详细信息如下："
    document.add_paragraph(temp_str)
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'名称'
    hdr_cells[1].text = u'更新时间'
    hdr_cells[2].text = u'官网网址'
    hdr_cells[3].text = u'备注信息'
    hdr_cells[4].text = u'城市'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[4].paragraphs[0].runs[0]
    run.font.bold = True

    for item in ixp_ca:
        # temp_str = str(item[0]), ",", str(item[1]), ",", str(item[2]), ",", str(item[3])
        # document.add_paragraph(temp_str)
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        row_cells[4].text = item[4]

    document.add_heading("- - - - - - -6)全球IXP的IPV6支持情况- - - - - - - - - - - -", level=1)
    temp_str = "已支持IPV6的IXP数量：", str(ipv6_on_cnt)
    document.add_paragraph(temp_str)
    temp_str = "未支持IPV6的IXP数量：", str(ipv6_off_cnt)
    document.add_paragraph(temp_str)

    document.add_heading("- - - - - - -7)案例1：阿姆斯特丹交换中心（AMS-IX）数据统计分析- - - - ", level=1)
    html = urlopen(r'https://peeringdb.com/api/ix/26')
    html_json = json.loads(html.read())
    # document.add_paragraph(html_json['data'][0])
    ix_data = html_json['data'][0]
    temp_str = "IXP名称(简称):", str(ix_data['name'])
    document.add_paragraph(temp_str)
    temp_str = "IXP名称(全称):", str(ix_data['name_long'])
    document.add_paragraph(temp_str)
    temp_str = "IXP所在城市及国家:", ix_data['city'], ", ", ix_data['country'], ", ", ix_data['region_continent']
    document.add_paragraph(temp_str)
    temp_str = "是否支持IPV6:", str(ix_data['proto_ipv6'])
    document.add_paragraph(temp_str)
    temp_str = "官方网站:", ix_data['website']
    document.add_paragraph(temp_str)
    temp_str = "该IXP流量信息展示页面:", ix_data['url_stats']
    document.add_paragraph(temp_str)
    temp_str = "该IXP接入网络数量:", str(ix_data['net_count'])
    document.add_paragraph(temp_str)
    document.add_paragraph("该IXP网络基础设施点:")
    item_cnt = 1
    for item in ix_data['fac_set']:
        temp_str = str(item_cnt), "> ", item['name'], ", ", item['city'], ", ", item['country']
        document.add_paragraph(temp_str)
        item_cnt += 1

    document.add_heading("- - - - - - -8)案例2：香港交换中心（HK-IX）数据统计分析- - ", level=1)
    html = urlopen(r'https://peeringdb.com/api/ix/42')
    html_json = json.loads(html.read())
    # document.add_paragraph(html_json['data'][0])
    ix_data = html_json['data'][0]
    temp_str = "IXP名称(简称):", str(ix_data['name'])
    document.add_paragraph(temp_str)
    temp_str = "IXP名称(全称):", str(ix_data['name_long'])
    document.add_paragraph(temp_str)
    temp_str = "IXP所在城市及国家:", ix_data['city'], ", ", ix_data['country'], ", ", ix_data['region_continent']
    document.add_paragraph(temp_str)
    temp_str = "是否支持IPV6:", str(ix_data['proto_ipv6'])
    document.add_paragraph(temp_str)
    temp_str = "官方网站:", ix_data['website']
    document.add_paragraph(temp_str)
    temp_str = "该IXP流量信息展示页面:", ix_data['url_stats']
    document.add_paragraph(temp_str)
    temp_str = "该IXP接入网络数量:", str(ix_data['net_count'])
    document.add_paragraph(temp_str)
    document.add_paragraph("该IXP网络基础设施点:")
    item_cnt = 1
    for item in ix_data['fac_set']:
        temp_str = str(item_cnt), "> ", item['name'], ", ", item['city'], ", ", item['country']
        document.add_paragraph(temp_str)
        item_cnt += 1

    document.add_heading("- - - - - - -9)案例3：莫斯科交换中心（MSK-IX Moscow）数据统计分析- -", level=1)
    html = urlopen(r'https://peeringdb.com/api/ix/100')
    html_json = json.loads(html.read())
    # document.add_paragraph(html_json['data'][0])
    ix_data = html_json['data'][0]
    temp_str = "IXP名称(简称):", str(ix_data['name'])
    document.add_paragraph(temp_str)
    temp_str = "IXP名称(全称):", str(ix_data['name_long'])
    document.add_paragraph(temp_str)
    temp_str = "IXP所在城市及国家:", ix_data['city'], ", ", ix_data['country'], ", ", ix_data['region_continent']
    document.add_paragraph(temp_str)
    temp_str = "是否支持IPV6:", str(ix_data['proto_ipv6'])
    document.add_paragraph(temp_str)
    temp_str = "官方网站:", ix_data['website']
    document.add_paragraph(temp_str)
    temp_str = "该IXP流量信息展示页面:", ix_data['url_stats']
    document.add_paragraph(temp_str)
    temp_str = "该IXP接入网络数量:", str(ix_data['net_count'])
    document.add_paragraph(temp_str)
    document.add_paragraph("该IXP网络基础设施点:")
    item_cnt = 1
    for item in ix_data['fac_set']:
        temp_str = str(item_cnt), "> ", item['name'], ", ", item['city'], ", ", item['country']
        document.add_paragraph(temp_str)
        item_cnt += 1

    document.add_page_break()
    document.save(doc_file_path)


if __name__ == "__main__":
    time_start = time.time()
    generate_global_ixp_report(r'./ixp-view.docx')
    time_end = time.time()
    print("=>Scripts Finish, Time Consuming:", (time_end - time_start), "S")
