# coding:utf-8
"""
create on Apr 19, 2021 By Wenyan YU
Email: ieeflsyu@outlook.com

测试自动化生成word报告的能力

"""
from docx import Document
from docx.shared import Inches


def gen_doc_file(doc_file_path):
    """
    生成word文档
    :param doc_file_path:
    :return None:
    """
    print(doc_file_path)
    # 新建一个文档
    document = Document()
    document.add_heading(u'自动分析报告', 0)

    document.add_heading(u'一、基本介绍', level=1)
    # 添加一个段落
    p = document.add_paragraph(u'python-doc模块是一个非常实用的用于自动生成报告的文档，可以自动根据读取的数据生成')
    p.add_run(u'图片').bold = True
    p.add_run(u'和')
    p.add_run(u'表格').italic = True

    document.add_paragraph(u'python-doc模块可以用于：')
    # 无序列表
    document.add_paragraph(
        u'根据程序计算动态结果替换动态内容，如统计数字等', style='List Bullet'
    )
    document.add_paragraph(
        u'可以自动嵌入相应的图片和表格', style='List Bullet'
    )
    document.add_paragraph(
        u'支持各类样式进行调整', style='List Bullet'
    )

    document.add_paragraph(u'python-doc模块不足的地方：')
    document.add_paragraph(
        u'相对简单', style='List Number'
    )
    document.add_paragraph(
        u'暂不支持WORD文档模板', style='List Number'
    )

    document.add_heading(u'二、板块统计', level=1)
    text = u'沪深两地的上市A股总共有3252只，其中沪市有1285只，深市有1967只,各板块的数据占比如下所示'
    document.add_paragraph(text)
    document.add_picture('./pic1.jpg', width=Inches(5.0))

    document.add_heading(u'三、时间统计', level=1)
    text = u'上市时间分布图如下所示，可以看出今明两年并不是上市的高峰期'
    document.add_paragraph(text)
    document.add_picture('./pic2.jpg', width=Inches(5.0))

    document.add_heading(u'四、新股统计', level=1)
    text = u'统计新股，列表如下：'
    document.add_paragraph(text)

    temp_list = [['海尔施', '上海', '主板'],
                 ['海尔施', '上海', '主板'],
                 ['海尔施', '上海', '主板'],
                 ['海尔施', '上海', '主板'],
                 ['海尔施', '上海', '主板']]

    # 插入表格
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'股票名称'
    hdr_cells[1].text = u'上市交易所'
    hdr_cells[2].text = u'上市板块'

    # 遍历表格
    for item in temp_list:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    document.add_page_break()
    document.save(doc_file_path)


if __name__ == "__main__":
    gen_doc_file(r'./test.docx')
